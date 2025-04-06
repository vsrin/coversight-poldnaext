import os
import tempfile
import json
import argparse
from docx import Document
from src.main import PolicyDNAExtractor
from config.config import get_config

def create_test_document():
    """Create a sample insurance policy document for testing."""
    doc = Document()
    
    # Add a title
    doc.add_heading('COMMERCIAL GENERAL LIABILITY POLICY', 0)
    
    # Add declarations section
    doc.add_heading('DECLARATIONS', 1)
    doc.add_paragraph('Policy Number: TEST-12345')
    doc.add_paragraph('Named Insured: ABC Corporation')
    doc.add_paragraph('Policy Period: 01/01/2025 to 01/01/2026')
    doc.add_paragraph('Limits of Insurance: $1,000,000 Each Occurrence')
    
    # Add insuring agreement
    doc.add_heading('SECTION I - INSURING AGREEMENT', 1)
    doc.add_paragraph(
        'We will pay those sums that the insured becomes legally obligated to pay as damages '
        'because of "bodily injury" or "property damage" to which this insurance applies. We will '
        'have the right and duty to defend the insured against any "suit" seeking those damages. '
        'However, we will have no duty to defend the insured against any "suit" seeking damages for '
        '"bodily injury" or "property damage" to which this insurance does not apply.'
    )
    
    # Add exclusions
    doc.add_heading('SECTION II - EXCLUSIONS', 1)
    doc.add_paragraph(
        'This insurance does not apply to:\n'
        '1. Expected or Intended Injury\n'
        '"Bodily injury" or "property damage" expected or intended from the standpoint of the insured.\n'
        '2. Contractual Liability\n'
        '"Bodily injury" or "property damage" for which the insured is obligated to pay damages by '
        'reason of the assumption of liability in a contract or agreement.'
    )
    
    # Add definitions
    doc.add_heading('SECTION III - DEFINITIONS', 1)
    doc.add_paragraph(
        '1. "Bodily injury" means bodily injury, sickness or disease sustained by a person, including '
        'death resulting from any of these at any time.\n'
        '2. "Property damage" means:\n'
        '   a. Physical injury to tangible property, including all resulting loss of use of that property; or\n'
        '   b. Loss of use of tangible property that is not physically injured.'
    )
    
    # Add conditions
    doc.add_heading('SECTION IV - CONDITIONS', 1)
    doc.add_paragraph(
        '1. Bankruptcy\n'
        'Bankruptcy or insolvency of the insured or of the insured\'s estate will not relieve us of our '
        'obligations under this policy.\n'
        '2. Duties In The Event Of Occurrence, Offense, Claim Or Suit\n'
        'a. You must see to it that we are notified as soon as practicable of an "occurrence" or an offense '
        'which may result in a claim.'
    )
    
    # Add endorsement
    doc.add_heading('ENDORSEMENT - ADDITIONAL INSURED', 1)
    doc.add_paragraph(
        'This endorsement modifies insurance provided under the following:\n'
        'COMMERCIAL GENERAL LIABILITY COVERAGE PART\n\n'
        'SCHEDULE\n'
        'Name of Additional Insured Person(s) Or Organization(s):\n'
        'XYZ Partner Company\n\n'
        'Section II – Who Is An Insured is amended to include as an additional insured the person(s) or '
        'organization(s) shown in the Schedule.'
    )
    
    # Save the document to a temporary file
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, 'sample_policy.docx')
    doc.save(file_path)
    
    print(f"Created sample policy document at: {file_path}")
    return file_path

def display_element_examples(document_map, element_type, count=3):
    """
    Display examples of specific element types from the document map.
    
    Args:
        document_map: The processed document map
        element_type: Type of element to display
        count: Number of examples to show
    """
    elements = [e for e in document_map.get('elements', []) if e.get('type') == element_type]
    
    if not elements:
        print(f"No {element_type} elements found.")
        return
    
    print(f"\nExamples of {element_type} elements:")
    for i, element in enumerate(elements[:count]):
        print(f"  Example {i+1}:")
        print(f"    Text: {element.get('text', '')[:150]}...")
        print(f"    Subtype: {element.get('subtype', 'None')}")
        print(f"    Confidence: {element.get('confidence', 0.0)}")
        if element.get('keywords'):
            print(f"    Keywords: {', '.join(element.get('keywords', []))}")
        print()

def display_language_analysis_examples(document_map, count=3):
    """
    Display examples of language analysis from the document map.
    
    Args:
        document_map: The processed document map with language analysis
        count: Number of examples to show
    """
    # Display intent analysis examples
    elements_with_intent = [e for e in document_map.get('elements_with_language_analysis', []) 
                           if e.get('intent_analysis', {}).get('intent_summary')]
    
    if elements_with_intent:
        print("\nExamples of Intent Analysis:")
        for i, element in enumerate(elements_with_intent[:count]):
            intent_analysis = element.get('intent_analysis', {})
            print(f"  Example {i+1}:")
            print(f"    Text: {element.get('text', '')[:100]}...")
            print(f"    Intent Summary: {intent_analysis.get('intent_summary', 'None')}")
            print(f"    Coverage Effect: {intent_analysis.get('coverage_effect', 'None')}")
            print(f"    Confidence: {intent_analysis.get('intent_confidence', 0.0)}")
            print()
    
    # Display conditional language examples
    elements_with_conditions = [e for e in document_map.get('elements_with_language_analysis', []) 
                               if e.get('conditional_analysis', {}).get('conditions')]
    
    if elements_with_conditions:
        print("\nExamples of Conditional Language:")
        for i, element in enumerate(elements_with_conditions[:count]):
            conditional_analysis = element.get('conditional_analysis', {})
            conditions = conditional_analysis.get('conditions', [])
            print(f"  Example {i+1}:")
            print(f"    Text: {element.get('text', '')[:100]}...")
            print(f"    Condition Count: {conditional_analysis.get('condition_count', 0)}")
            if conditions:
                print(f"    First Condition: {conditions[0].get('condition_text', 'None')}")
                print(f"    Condition Type: {conditions[0].get('condition_type', 'None')}")
            print()
    
    # Display term extraction examples
    elements_with_terms = [e for e in document_map.get('elements_with_language_analysis', []) 
                          if e.get('term_extraction', {}).get('extracted_terms')]
    
    if elements_with_terms:
        print("\nExamples of Term Extraction:")
        for i, element in enumerate(elements_with_terms[:count]):
            term_extraction = element.get('term_extraction', {})
            extracted_terms = term_extraction.get('extracted_terms', [])
            print(f"  Example {i+1}:")
            print(f"    Text: {element.get('text', '')[:100]}...")
            print(f"    Extracted Terms: {len(extracted_terms)}")
            if extracted_terms:
                print(f"    First Term: {extracted_terms[0].get('term', 'None')}")
                print(f"    Term Type: {extracted_terms[0].get('term_type', 'None')}")
            print()

def display_relationship_examples(document_map, count=3):
    """
    Display examples of cross-references and dependencies from the document map.
    
    Args:
        document_map: The processed document map with cross-reference data
        count: Number of examples to show
    """
    # Check if cross-reference map exists
    if 'cross_reference_map' not in document_map:
        print("No cross-reference data found.")
        return
    
    cross_ref_map = document_map['cross_reference_map']
    
    # Display reference examples
    edges = cross_ref_map.get('edges', [])
    reference_edges = [e for e in edges if e.get('type') == 'reference']
    
    if reference_edges:
        print("\nExamples of Cross-References:")
        for i, edge in enumerate(reference_edges[:count]):
            print(f"  Example {i+1}:")
            print(f"    Type: {edge.get('subtype', 'Unknown')}")
            print(f"    Text: {edge.get('text', 'None')}")
            print(f"    Confidence: {edge.get('weight', 0.0)}")
            print()
    
    # Display dependency examples
    dependency_edges = [e for e in edges if e.get('type') == 'dependency']
    
    if dependency_edges:
        print("\nExamples of Dependencies:")
        for i, edge in enumerate(dependency_edges[:count]):
            print(f"  Example {i+1}:")
            print(f"    Type: {edge.get('subtype', 'Unknown')}")
            print(f"    Strength: {edge.get('weight', 0.0)}")
            print(f"    Origin: {edge.get('metadata', {}).get('origin', 'Unknown')}")
            print()
    
    # Display conflict examples
    conflicts = cross_ref_map.get('conflicts', [])
    
    if conflicts:
        print("\nExamples of Conflicts:")
        for i, conflict in enumerate(conflicts[:count]):
            print(f"  Example {i+1}:")
            print(f"    Type: {conflict.get('conflict_type', 'Unknown')}")
            print(f"    Description: {conflict.get('description', 'None')}")
            print(f"    Severity: {conflict.get('severity', 0.0)}")
            print()

# New function to display taxonomy examples
def display_taxonomy_examples(document_map, count=3):
    """
    Display examples of taxonomy mappings and normalized language from the document map.
    
    Args:
        document_map: The processed document map with taxonomy data
        count: Number of examples to show
    """
    # Check if standardized policy structure exists
    if 'standardized_policy_structure' not in document_map:
        print("No taxonomy standardization data found.")
        return
    
    policy_structure = document_map['standardized_policy_structure']
    
    # Display taxonomy mapping examples
    taxonomy_mappings = policy_structure.get('taxonomy_mappings', {})
    if taxonomy_mappings:
        print("\nExamples of Taxonomy Mappings:")
        mapped_elements = []
        
        # Get elements with mappings
        elements = policy_structure.get('elements', {})
        for element_id, mapping in list(taxonomy_mappings.items())[:count]:
            if element_id in elements:
                element = elements[element_id]
                primary_mapping = mapping.get('primary_mapping', {})
                mapped_elements.append({
                    'element': element,
                    'mapping': primary_mapping
                })
        
        # Display mapped elements
        for i, item in enumerate(mapped_elements):
            element = item['element']
            mapping = item['mapping']
            print(f"  Example {i+1}:")
            print(f"    Text: {element.get('text', '')[:100]}...")
            print(f"    Type: {element.get('type', 'Unknown')}")
            print(f"    Taxonomy Code: {mapping.get('code', 'Unknown')}")
            print(f"    Confidence: {mapping.get('confidence', 0.0):.2f}")
            print()
    
    # Display normalized language examples
    normalized_language = policy_structure.get('normalized_language', {})
    if normalized_language:
        print("\nExamples of Normalized Language:")
        for i, (element_id, norm_info) in enumerate(list(normalized_language.items())[:count]):
            print(f"  Example {i+1}:")
            if element_id in elements:
                element = elements[element_id]
                print(f"    Original Text: {element.get('text', '')[:80]}...")
                print(f"    Normalized Text: {norm_info.get('normalized_text', '')[:80]}...")
                print(f"    Source: {norm_info.get('normalization_source', 'Unknown')}")
                print(f"    Is Unique: {norm_info.get('uniqueness_analysis', {}).get('is_unique', False)}")
                print()
    
    # Display unique provisions if available
    unique_provisions = []
    for element_id, norm_info in normalized_language.items():
        uniqueness_analysis = norm_info.get('uniqueness_analysis', {})
        if uniqueness_analysis.get('is_unique', False) and element_id in elements:
            unique_provisions.append({
                'element': elements[element_id],
                'uniqueness_score': uniqueness_analysis.get('uniqueness_score', 0.0),
                'unique_phrases': uniqueness_analysis.get('unique_phrases', [])
            })
    
    if unique_provisions:
        print("\nExamples of Unique Provisions:")
        for i, item in enumerate(unique_provisions[:count]):
            element = item['element']
            print(f"  Example {i+1}:")
            print(f"    Text: {element.get('text', '')[:100]}...")
            print(f"    Type: {element.get('type', 'Unknown')}")
            print(f"    Uniqueness Score: {item['uniqueness_score']:.2f}")
            if item['unique_phrases']:
                print(f"    First Unique Phrase: {item['unique_phrases'][0][:80]}...")
            print()

def main():
    """Main function to demonstrate the Policy DNA Extractor with modular phase execution."""
    # Set up command-line argument parsing
    parser = argparse.ArgumentParser(description="Demonstrate the Policy DNA Extractor")
    parser.add_argument("--phase", type=int, choices=[1, 2, 3, 4, 5], 
                        help="Processing phase to run: 1=Document Processing, 2=Element Extraction, 3=Deep Language Analysis, 4=Cross-Reference Mapping, 5=Taxonomy Standardization")
    parser.add_argument("--input", type=str, help="Path to input document or previous phase result")
    parser.add_argument("--output-dir", type=str, default="output", help="Directory to save output files")
    args = parser.parse_args()
    
    # Default to running all phases if not specified
    phase = args.phase if args.phase else 5
    
    print("Policy DNA Extractor Example")
    print("===========================")
    
    # Initialize configuration
    config = get_config()
    config.output_dir = args.output_dir if args.output_dir else os.path.join(os.getcwd(), 'output')
    os.makedirs(config.output_dir, exist_ok=True)
    config.debug_mode = True  # Enable debug mode for this demo
    
    # Create the extractor
    extractor = PolicyDNAExtractor(config)
    
    # Determine input document or previous phase result
    document_path = None
    previous_result = None
    
    if args.input:
        if args.input.endswith('.pdf') or args.input.endswith('.docx'):
            document_path = args.input
        elif args.input.endswith('.json'):
            try:
                with open(args.input, 'r') as f:
                    previous_result = json.load(f)
                print(f"Loaded input from: {args.input}")
            except FileNotFoundError:
                print(f"Error: Input file not found: {args.input}")
                return
            except json.JSONDecodeError:
                print(f"Error: Invalid JSON in input file: {args.input}")
                return
        else:
            print(f"Unsupported input file: {args.input}")
            return
    else:
        # Create sample document if no input provided
        print("\nStep 0: Creating a sample policy document...")
        document_path = create_test_document()
    
    result = None
    classified_sections = None
    
    # Phase 1: Document Processing and Segmentation
    if phase == 1 or phase == 0:
        print("\nPhase 1: Document Processing and Segmentation")
        print("--------------------------------------------")
        
        if document_path:
            # Parse document
            print("Step 1.1: Parsing document...")
            document_info = extractor.document_parser.parse_document(document_path)
            
            # Analyze structure
            print("Step 1.2: Analyzing document structure...")
            document_structure = extractor.structure_analyzer.analyze_structure(document_info)
            
            # Classify sections
            print("Step 1.3: Classifying sections...")
            classified_sections = extractor.section_classifier.classify_sections(document_structure['sections'])
            
            # Create document map
            print("Step 1.4: Creating document map...")
            document_map = extractor.document_mapper.create_document_map(document_info, classified_sections)
            
            # Save phase 1 results
            phase1_output = os.path.join(config.output_dir, 'phase1_document_map.json')
            with open(phase1_output, 'w') as f:
                json.dump(document_map, f, indent=2)
            
            print(f"Phase 1 results saved to: {phase1_output}")
            result = document_map
            
            # Display summary
            print("\nPhase 1 Summary:")
            print(f"Document ID: {document_map['document_id']}")
            print(f"Sections found: {document_map['section_counts']['TOTAL']}")
            
            print("\nSection count by type:")
            for section_type, count in sorted(document_map['section_counts'].items()):
                if section_type != 'TOTAL':
                    print(f"  - {section_type}: {count}")
    
    # Phase 2: Element Extraction and Classification
    if phase == 2 or phase == 0:
        print("\nPhase 2: Element Extraction and Classification")
        print("--------------------------------------------")
        
        # Use previous result or load phase 1 result
        if phase != 1 and not result:
            if previous_result:
                document_map = previous_result
                classified_sections = document_map.get('sections', [])
                print("Using provided input file as Phase 1 result.")
            else:
                phase1_output = os.path.join(config.output_dir, 'phase1_document_map.json')
                if os.path.exists(phase1_output):
                    with open(phase1_output, 'r') as f:
                        document_map = json.load(f)
                    classified_sections = document_map.get('sections', [])
                    print(f"Loaded Phase 1 results from: {phase1_output}")
                else:
                    print("Error: Phase 1 results not found. Please run Phase 1 first.")
                    return
        
        # Extract elements
        print("Step 2.1: Extracting policy elements...")
        all_elements = []
        
        for section in classified_sections:
            try:
                print(f"  Extracting elements from section: {section.get('title', 'Untitled')}")
                elements = extractor.element_extractor.extract_elements(section)
                
                if elements:
                    print(f"  Found {len(elements)} elements")
                    all_elements.extend(elements)
            except Exception as e:
                print(f"  Error extracting elements from section {section.get('title', 'Untitled')}: {str(e)}")
        
        # Classify elements
        print("Step 2.2: Classifying policy elements...")
        classified_elements = []
        
        for section in classified_sections:
            try:
                # Get elements for this section
                section_elements = [e for e in all_elements if e.get('section_id') == section.get('id')]
                
                if section_elements:
                    print(f"  Classifying elements in section: {section.get('title', 'Untitled')}")
                    classified = extractor.element_classifier.classify_elements(section_elements, section)
                    classified_elements.extend(classified)
            except Exception as e:
                print(f"  Error classifying elements in section {section.get('title', 'Untitled')}: {str(e)}")
        
        # Analyze relationships
        print("Step 2.3: Analyzing element relationships...")
        enhanced_elements = []
        
        for section in classified_sections:
            try:
                # Get elements for this section
                section_elements = [e for e in classified_elements if e.get('section_id') == section.get('id')]
                
                if section_elements:
                    print(f"  Analyzing relationships in section: {section.get('title', 'Untitled')}")
                    with_relationships = extractor.relationship_analyzer.analyze_relationships(section_elements, section)
                    enhanced_elements.extend(with_relationships)
            except Exception as e:
                print(f"  Error analyzing relationships in section {section.get('title', 'Untitled')}: {str(e)}")
                section_elements = [e for e in classified_elements if e.get('section_id') == section.get('id')]
                enhanced_elements.extend(section_elements)
        
        # Create element map
        print("Step 2.4: Creating element map...")
        enhanced_document_map = extractor.element_mapper.create_element_map(enhanced_elements, document_map)
        
        # Save phase 2 results
        phase2_output = os.path.join(config.output_dir, 'phase2_element_map.json')
        with open(phase2_output, 'w') as f:
            json.dump(enhanced_document_map, f, indent=2)
        
        print(f"Phase 2 results saved to: {phase2_output}")
        result = enhanced_document_map
        
        # Display summary
        print("\nPhase 2 Summary:")
        print(f"Elements extracted: {enhanced_document_map['element_counts']['TOTAL']}")
        
        print("\nElement count by type:")
        for element_type, count in sorted(enhanced_document_map['element_counts'].items()):
            if element_type != 'TOTAL':
                print(f"  - {element_type}: {count}")
        
        # Display examples of different element types
        display_element_examples(enhanced_document_map, "COVERAGE_GRANT")
        display_element_examples(enhanced_document_map, "EXCLUSION")
    
    # Phase 3: Deep Language Analysis
    if phase == 3 or phase == 0:
        print("\nPhase 3: Deep Language Analysis")
        print("-----------------------------")
        
        # Use previous result or load phase 2 result if not continuing from phase 2
        if phase != 2 and not result:
            if previous_result:
                enhanced_document_map = previous_result
                enhanced_elements = enhanced_document_map.get('elements', [])
                print("Using provided input file as Phase 2 result.")
            else:
                phase2_output = os.path.join(config.output_dir, 'phase2_element_map.json')
                if os.path.exists(phase2_output):
                    with open(phase2_output, 'r') as f:
                        enhanced_document_map = json.load(f)
                    enhanced_elements = enhanced_document_map.get('elements', [])
                    print(f"Loaded Phase 2 results from: {phase2_output}")
                else:
                    print("Error: Phase 2 results not found. Please run Phase 2 first.")
                    return
        else:
            enhanced_elements = result.get('elements', [])
        
        # Step 3.1: Analyze element intent
        print("Step 3.1: Analyzing element intent...")
        try:
            elements_with_intent = extractor.intent_analyzer.analyze_intent(enhanced_elements)
            print(f"  Analyzed intent for {len(elements_with_intent)} elements")
            
            # Save intermediate results if in debug mode
            intent_output = os.path.join(config.output_dir, 'debug', 'elements_with_intent.json')
            os.makedirs(os.path.dirname(intent_output), exist_ok=True)
            with open(intent_output, 'w') as f:
                json.dump(elements_with_intent, f, indent=2)
        except Exception as e:
            print(f"  Error analyzing element intent: {str(e)}")
            elements_with_intent = enhanced_elements  # Fall back to previous elements
        
        # Step 3.2: Detect conditional language
        print("Step 3.2: Detecting conditional language...")
        try:
            elements_with_conditions = extractor.conditional_language_detector.detect_conditions(elements_with_intent)
            print(f"  Detected conditions in elements")
            
            # Save intermediate results if in debug mode
            conditions_output = os.path.join(config.output_dir, 'debug', 'elements_with_conditions.json')
            with open(conditions_output, 'w') as f:
                json.dump(elements_with_conditions, f, indent=2)
        except Exception as e:
            print(f"  Error detecting conditional language: {str(e)}")
            elements_with_conditions = elements_with_intent  # Fall back to previous elements
        
        # Step 3.3: Extract specific terms
        print("Step 3.3: Extracting specific terms...")
        try:
            elements_with_terms = extractor.term_extractor.extract_terms(elements_with_conditions)
            print(f"  Extracted terms from elements")
            
            # Save intermediate results if in debug mode
            terms_output = os.path.join(config.output_dir, 'debug', 'elements_with_terms.json')
            with open(terms_output, 'w') as f:
                json.dump(elements_with_terms, f, indent=2)
        except Exception as e:
            print(f"  Error extracting specific terms: {str(e)}")
            elements_with_terms = elements_with_conditions  # Fall back to previous elements
        
        # Step 3.4: Create language map
        print("Step 3.4: Creating language map...")
        try:
            final_document_map = extractor.language_mapper.create_language_map(elements_with_terms, enhanced_document_map)
            print(f"  Created language map successfully")
            
            # Save phase 3 results
            phase3_output = os.path.join(config.output_dir, 'phase3_language_map.json')
            with open(phase3_output, 'w') as f:
                json.dump(final_document_map, f, indent=2)
            
            print(f"Phase 3 results saved to: {phase3_output}")
            result = final_document_map
            
            # Display language analysis examples
            display_language_analysis_examples(final_document_map)
            
            # Display language insights
            if 'language_insights' in final_document_map:
                insights = final_document_map['language_insights']
                
                print("\nLanguage Analysis Insights:")
                
                # Coverage summary
                coverage_summary = insights.get('coverage_summary', {})
                print(f"  Coverage grants: {len(coverage_summary.get('coverage_grants', []))}")
                print(f"  Key exclusions: {len(coverage_summary.get('key_exclusions', []))}")
                print(f"  Key limitations: {len(coverage_summary.get('key_limitations', []))}")
                
                # Conditions
                key_conditions = insights.get('key_conditions', [])
                print(f"  Key conditions: {len(key_conditions)}")
                
                # Defined terms
                defined_terms = insights.get('defined_terms_usage', {})
                print(f"  Defined terms: {defined_terms.get('defined_terms_count', 0)}")
                print(f"  Terms with definitions: {defined_terms.get('terms_with_definitions', 0)}")
                
                # Interpretation challenges
                challenges = insights.get('interpretation_challenges', [])
                print(f"  Potential interpretation challenges: {len(challenges)}")
                
        except Exception as e:
            print(f"  Error creating language map: {str(e)}")
            # If unable to create final map, save the elements with terms as a fallback
            fallback_output = os.path.join(config.output_dir, 'phase3_fallback.json')
            with open(fallback_output, 'w') as f:
                json.dump(elements_with_terms, f, indent=2)
            print(f"  Saved fallback results to: {fallback_output}")

            # Phase 4: Cross-Reference and Dependency Mapping
    if phase == 4 or phase == 0:
        print("\nPhase 4: Cross-Reference and Dependency Mapping")
        print("---------------------------------------------")
        
        # Load Phase 3 result if not continuing from Phase 3
        if phase != 3 and not result:
            if previous_result:
                final_document_map = previous_result
                print("Using provided input file as Phase 3 result.")
            else:
                phase3_output = os.path.join(config.output_dir, 'phase3_language_map.json')
                if os.path.exists(phase3_output):
                    with open(phase3_output, 'r') as f:
                        final_document_map = json.load(f)
                    print(f"Loaded Phase 3 results from: {phase3_output}")
                else:
                    print("Error: Phase 3 results not found. Please run Phase 3 first.")
                    return
        else:
            final_document_map = result
        
        # Step 4.1: Detect references
        print("Step 4.1: Detecting cross-references...")
        try:
            references = extractor.reference_detector.detect_references(final_document_map)
            print(f"  Detected {references.get('total_references', 0)} references")
            
            # Save intermediate results if in debug mode
            references_output = os.path.join(config.output_dir, 'debug', 'references.json')
            os.makedirs(os.path.dirname(references_output), exist_ok=True)
            with open(references_output, 'w') as f:
                json.dump(references, f, indent=2)
        except Exception as e:
            print(f"  Error detecting references: {str(e)}")
            references = {"references": [], "reference_type_counts": {}, "total_references": 0}
        
        # Step 4.2: Analyze dependencies
        print("Step 4.2: Analyzing logical dependencies...")
        try:
            dependencies = extractor.dependency_analyzer.analyze_dependencies(final_document_map, references)
            print(f"  Identified {dependencies.get('total_dependencies', 0)} dependencies")
            
            # Save intermediate results if in debug mode
            dependencies_output = os.path.join(config.output_dir, 'debug', 'dependencies.json')
            with open(dependencies_output, 'w') as f:
                json.dump(dependencies, f, indent=2)
        except Exception as e:
            print(f"  Error analyzing dependencies: {str(e)}")
            dependencies = {"dependencies": [], "dependency_type_counts": {}, "total_dependencies": 0}
        
        # Step 4.3: Identify conflicts
        print("Step 4.3: Identifying potential conflicts...")
        try:
            conflicts = extractor.conflict_identifier.identify_conflicts(final_document_map, dependencies)
            print(f"  Found {conflicts.get('total_conflicts', 0)} potential conflicts")
            
            # Save intermediate results if in debug mode
            conflicts_output = os.path.join(config.output_dir, 'debug', 'conflicts.json')
            with open(conflicts_output, 'w') as f:
                json.dump(conflicts, f, indent=2)
        except Exception as e:
            print(f"  Error identifying conflicts: {str(e)}")
            conflicts = {"conflicts": [], "conflict_type_counts": {}, "total_conflicts": 0}
        
        # Step 4.4: Build relationship graph
        print("Step 4.4: Building relationship graph...")
        try:
            graph_result = extractor.graph_builder.build_graph(final_document_map, references, dependencies, conflicts)
            final_document_map['cross_reference_map'] = graph_result
            print(f"  Built graph with {graph_result['graph_stats']['node_count']} nodes and {graph_result['graph_stats']['edge_count']} edges")
            
            # Save intermediate results if in debug mode
            graph_output = os.path.join(config.output_dir, 'debug', 'relationship_graph.json')
            with open(graph_output, 'w') as f:
                json.dump(graph_result, f, indent=2)
        except Exception as e:
            print(f"  Error building relationship graph: {str(e)}")
            graph_result = {
                "nodes": [],
                "edges": [],
                "graph_stats": {"node_count": 0, "edge_count": 0}
            }
            final_document_map['cross_reference_map'] = graph_result
        
        # Save phase 4 results
        phase4_output = os.path.join(config.output_dir, 'phase4_graph_map.json')
        with open(phase4_output, 'w') as f:
            json.dump(final_document_map, f, indent=2)
        
        print(f"Phase 4 results saved to: {phase4_output}")
        result = final_document_map
        
        # Display relationship examples
        display_relationship_examples(final_document_map)
        
        # Display graph summary
        print("\nPhase 4 Summary:")
        print(f"  Total references detected: {references.get('total_references', 0)}")
        print(f"  Total dependencies identified: {dependencies.get('total_dependencies', 0)}")
        print(f"  Potential conflicts found: {conflicts.get('total_conflicts', 0)}")
        print(f"  Graph nodes: {graph_result['graph_stats']['node_count']}")
        print(f"  Graph edges: {graph_result['graph_stats']['edge_count']}")
        
        # Display connectivity statistics
        if 'connectivity' in graph_result['graph_stats']:
            conn = graph_result['graph_stats']['connectivity']
            print("\nGraph Connectivity:")
            print(f"  Connected components: {conn.get('connected_components', 0)}")
            print(f"  Largest component size: {conn.get('largest_component_size', 0)} nodes")
            print(f"  Isolated nodes: {conn.get('isolated_nodes', 0)} ({conn.get('isolated_percentage', 0)}%)")
        
        # Display top referenced elements
        if 'most_referenced' in graph_result and graph_result['most_referenced']:
            print("\nMost Referenced Elements:")
            for i, element in enumerate(graph_result['most_referenced'][:3]):
                print(f"  {i+1}. {element.get('element_text', '')[:50]}... ({element.get('reference_count', 0)} references)")
        
        # Display reference types
        if 'reference_type_counts' in references:
            print("\nReference Type Counts:")
            for ref_type, count in references['reference_type_counts'].items():
                print(f"  - {ref_type}: {count}")
        
        # Display conflict types if any found
        if conflicts.get('total_conflicts', 0) > 0 and 'conflict_type_counts' in conflicts:
            print("\nConflict Type Counts:")
            for conflict_type, count in conflicts['conflict_type_counts'].items():
                print(f"  - {conflict_type}: {count}")
    
    # Phase 5: Standardization and Taxonomy Mapping (New)
    if phase == 5 or phase == 0:
        print("\nPhase 5: Standardization and Taxonomy Mapping")
        print("------------------------------------------")
        
        # Load Phase 4 result if not continuing from Phase 4
        if phase != 4 and not result:
            if previous_result:
                final_document_map = previous_result
                print("Using provided input file as Phase 4 result.")
            else:
                phase4_output = os.path.join(config.output_dir, 'phase4_graph_map.json')
                if os.path.exists(phase4_output):
                    with open(phase4_output, 'r') as f:
                        final_document_map = json.load(f)
                    print(f"Loaded Phase 4 results from: {phase4_output}")
                else:
                    print("Error: Phase 4 results not found. Please run Phase 4 first.")
                    return
        else:
            # Use the result from Phase 4
            final_document_map = result
        
        # Extract elements for mapping
        elements = final_document_map.get('elements', [])
        
        # Step 5.1: Map elements to standardized taxonomy
        print("Step 5.1: Mapping elements to standardized taxonomy...")
        try:
            taxonomy_mapping_results = extractor.taxonomy_mapper.map_elements(elements)
            
            # Convert taxonomy mapping results to dictionary for JSON serialization
            taxonomy_mappings_dict = {
                element_id: mapping_result.to_dict() 
                for element_id, mapping_result in taxonomy_mapping_results.items()
            }
            
            # Add taxonomy mapping statistics
            taxonomy_stats = extractor.taxonomy_mapper.get_confidence_statistics(taxonomy_mapping_results)
            taxonomy_distribution = extractor.taxonomy_mapper.get_taxonomy_distribution(taxonomy_mapping_results)
            
            final_document_map['taxonomy_mapping_stats'] = taxonomy_stats
            final_document_map['taxonomy_distribution'] = taxonomy_distribution
            
            # Save intermediate results if in debug mode
            mappings_output = os.path.join(config.output_dir, 'debug', 'taxonomy_mappings.json')
            os.makedirs(os.path.dirname(mappings_output), exist_ok=True)
            with open(mappings_output, 'w') as f:
                json.dump(taxonomy_mappings_dict, f, indent=2)
                
            print(f"  Mapped {len(taxonomy_mapping_results)} elements to standardized taxonomy")
            print(f"  Average mapping confidence: {taxonomy_stats['avg_confidence']:.2f}")
            print(f"  High confidence mappings: {taxonomy_stats['high_confidence_count']}")
            
        except Exception as e:
            print(f"  Error mapping elements to taxonomy: {str(e)}")
            taxonomy_mapping_results = {}
            taxonomy_mappings_dict = {}
            final_document_map['taxonomy_mapping_stats'] = {
                "avg_confidence": 0,
                "min_confidence": 0,
                "max_confidence": 0,
                "high_confidence_count": 0,
                "medium_confidence_count": 0,
                "low_confidence_count": 0
            }
            final_document_map['taxonomy_distribution'] = {}
        
        # Step 5.2: Normalize policy language
        print("Step 5.2: Normalizing policy language...")
        try:
            normalized_elements = extractor.language_normalizer.normalize_elements(elements)
            
            # Generate normalization report
            normalization_report = extractor.language_normalizer.generate_normalization_report(normalized_elements)
            final_document_map['language_normalization_report'] = normalization_report
            
            # Save intermediate results if in debug mode
            normalized_output = os.path.join(config.output_dir, 'debug', 'normalized_elements.json')
            with open(normalized_output, 'w') as f:
                json.dump(normalized_elements, f, indent=2)
                
            print(f"  Normalized {len(normalized_elements)} policy elements")
            print(f"  Standardized elements: {normalization_report['standardized_count']} ({normalization_report['standardized_percentage']:.1f}%)")
            print(f"  Unique provisions: {normalization_report['unique_count']} ({normalization_report['unique_percentage']:.1f}%)")
            
        except Exception as e:
            print(f"  Error normalizing policy language: {str(e)}")
            normalized_elements = elements
            final_document_map['language_normalization_report'] = {
                "total_elements": len(elements),
                "standardized_count": 0,
                "standardized_percentage": 0,
                "unique_count": 0,
                "unique_percentage": 0,
                "average_similarity_score": 0
            }
        
        # Step 5.3: Build structured policy representation
        print("Step 5.3: Building structured policy representation...")
        try:
            # Extract policy metadata
            metadata = extractor._extract_policy_metadata(elements, final_document_map.get('document_info', {}))
            
            # Build the structured policy representation
            extractor.policy_structure_builder.set_policy_metadata(metadata)
            extractor.policy_structure_builder.set_document_map(final_document_map)
            extractor.policy_structure_builder.add_elements(elements)
            extractor.policy_structure_builder.add_taxonomy_mappings(taxonomy_mappings_dict)
            extractor.policy_structure_builder.add_normalized_language(normalized_elements)
            
            # Add relationships from Phase 4
            if 'cross_reference_map' in final_document_map:
                cross_ref_map = final_document_map['cross_reference_map']
                relationships = []
                
                # Convert graph edges to relationships
                for edge in cross_ref_map.get('edges', []):
                    relationship = {
                        "source_id": edge.get('source'),
                        "target_id": edge.get('target'),
                        "type": edge.get('type'),
                        "subtype": edge.get('subtype', ''),
                        "weight": edge.get('weight', 0)
                    }
                    relationships.append(relationship)
                
                extractor.policy_structure_builder.add_relationships(relationships)
            
            # Build the final policy structure
            policy_structure = extractor.policy_structure_builder.build_structure()
            final_document_map['standardized_policy_structure'] = policy_structure
            
            # Generate coverage summary
            coverage_summary = extractor.policy_structure_builder.get_coverage_summary()
            final_document_map['standardized_coverage_summary'] = coverage_summary
            
            # Save intermediate results if in debug mode
            structure_output = os.path.join(config.output_dir, 'debug', 'policy_structure.json')
            with open(structure_output, 'w') as f:
                json.dump(policy_structure, f, indent=2)
                
            coverage_output = os.path.join(config.output_dir, 'debug', 'coverage_summary.json')
            with open(coverage_output, 'w') as f:
                json.dump(coverage_summary, f, indent=2)
                
            print(f"  Created structured representation with {policy_structure['summary']['total_elements']} elements")
            print(f"  Mapped to {len(policy_structure['summary'].get('taxonomy_codes', {}))} taxonomy categories")
            
        except Exception as e:
            print(f"  Error building structured policy representation: {str(e)}")
            final_document_map['standardized_policy_structure'] = {"error": "Failed to build structured representation"}
            final_document_map['standardized_coverage_summary'] = {}
        
        # Step 5.4: Generate taxonomy visualizations
        print("Step 5.4: Generating taxonomy visualizations...")
        try:
            # Create visualizations directory
            vis_dir = os.path.join(config.output_dir, "visualizations")
            os.makedirs(vis_dir, exist_ok=True)
            
            # Generate base filename
            base_name = os.path.basename(document_path) if document_path else "policy"
            file_name = os.path.splitext(base_name)[0]
            
            # Check if we have a valid policy structure
            if 'standardized_policy_structure' in final_document_map and 'error' not in final_document_map['standardized_policy_structure']:
                policy_structure = final_document_map['standardized_policy_structure']
                
                # Generate HTML tree visualization
                tree_path = os.path.join(vis_dir, f"{file_name}_taxonomy_tree.html")
                extractor.taxonomy_visualizer.generate_html_tree(policy_structure, tree_path)
                
                # Generate coverage report
                coverage_path = os.path.join(vis_dir, f"{file_name}_coverage_report.html")
                extractor.taxonomy_visualizer.generate_coverage_report(policy_structure, coverage_path)
                
                # Generate uniqueness report
                uniqueness_path = os.path.join(vis_dir, f"{file_name}_uniqueness_report.html")
                extractor.taxonomy_visualizer.generate_uniqueness_report(policy_structure, uniqueness_path)
                
                # Generate JSON for external visualization
                json_path = os.path.join(vis_dir, f"{file_name}_visualization_data.json")
                extractor.taxonomy_visualizer.generate_json_visualization(policy_structure, json_path)
                
                # Add visualization paths to document map
                final_document_map['taxonomy_visualizations'] = {
                    "taxonomy_tree": tree_path,
                    "coverage_report": coverage_path,
                    "uniqueness_report": uniqueness_path,
                    "visualization_data": json_path
                }
                
                print(f"  Generated taxonomy visualizations in: {vis_dir}")
                print(f"  - Taxonomy Tree: {os.path.basename(tree_path)}")
                print(f"  - Coverage Report: {os.path.basename(coverage_path)}")
                print(f"  - Uniqueness Report: {os.path.basename(uniqueness_path)}")
                print(f"  - Visualization Data: {os.path.basename(json_path)}")
                
            else:
                print("  Skipping visualization generation due to missing policy structure")
                
        except Exception as e:
            print(f"  Error generating taxonomy visualizations: {str(e)}")
            final_document_map['taxonomy_visualizations'] = {}
        
        # Save phase 5 results
        phase5_output = os.path.join(config.output_dir, 'phase5_taxonomy_map.json')
        with open(phase5_output, 'w') as f:
            json.dump(final_document_map, f, indent=2)
        
        print(f"Phase 5 results saved to: {phase5_output}")
        result = final_document_map
        
        # Save final complete results
        final_output = os.path.join(config.output_dir, 'policy_dna_complete.json')
        with open(final_output, 'w') as f:
            json.dump(final_document_map, f, indent=2)
        
        print(f"Complete Policy DNA saved to: {final_output}")
        
        # Display taxonomy examples
        display_taxonomy_examples(final_document_map)
        
        # Display taxonomy summary
        print("\nPhase 5 Summary:")
        taxonomy_stats = final_document_map.get('taxonomy_mapping_stats', {})
        print(f"  Average mapping confidence: {taxonomy_stats.get('avg_confidence', 0):.2f}")
        print(f"  High confidence mappings: {taxonomy_stats.get('high_confidence_count', 0)}")
        print(f"  Medium confidence mappings: {taxonomy_stats.get('medium_confidence_count', 0)}")
        print(f"  Low confidence mappings: {taxonomy_stats.get('low_confidence_count', 0)}")
        
        norm_report = final_document_map.get('language_normalization_report', {})
        print(f"  Standardized elements: {norm_report.get('standardized_count', 0)} ({norm_report.get('standardized_percentage', 0):.1f}%)")
        print(f"  Unique provisions: {norm_report.get('unique_count', 0)} ({norm_report.get('unique_percentage', 0):.1f}%)")
        
        # Display taxonomy distribution
        taxonomy_dist = final_document_map.get('taxonomy_distribution', {})
        if taxonomy_dist:
            print("\nTop Taxonomy Categories:")
            sorted_dist = sorted(taxonomy_dist.items(), key=lambda x: x[1], reverse=True)
            for i, (code, count) in enumerate(sorted_dist[:5]):
                print(f"  {i+1}. {code}: {count} elements")
    
    print("\nProcessing completed successfully!")

if __name__ == "__main__":
    main()
    