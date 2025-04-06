"""
Document parser module for extracting text from insurance policy documents.
"""

import os
import io
from typing import Dict, List, Optional
import PyPDF2
from docx import Document
from config.config import ParserConfig

class DocumentParser:
    """Handles document file parsing and text extraction."""
    
    def __init__(self, config: Optional[ParserConfig] = None):
        """
        Initialize the document parser.
        
        Args:
            config: Configuration for the parser
        """
        self.config = config or ParserConfig()
    
    def parse_document(self, document_path: str) -> Dict:
        """
        Parse a document based on its file type.
        
        Args:
            document_path: Path to the document file
            
        Returns:
            Dict containing document text and metadata
            
        Raises:
            ValueError: If the file type is not supported
            FileNotFoundError: If the file does not exist
        """
        if not os.path.exists(document_path):
            raise FileNotFoundError(f"Document not found: {document_path}")
        
        file_extension = document_path.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return self._parse_pdf(document_path)
        elif file_extension in ['docx', 'doc']:
            return self._parse_docx(document_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def _parse_pdf(self, pdf_path: str) -> Dict:
        """
        Extract text and structure from PDF documents.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dict containing document text and metadata
        """
        document_info = {
            'file_path': pdf_path,
            'file_type': 'pdf',
            'pages': [],
            'full_text': "",
            'metadata': {}
        }
        
        try:
            with open(pdf_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if reader.metadata:
                    document_info['metadata'] = {
                        'title': reader.metadata.get('/Title', ''),
                        'author': reader.metadata.get('/Author', ''),
                        'subject': reader.metadata.get('/Subject', ''),
                        'creator': reader.metadata.get('/Creator', ''),
                        'producer': reader.metadata.get('/Producer', ''),
                        'creation_date': reader.metadata.get('/CreationDate', '')
                    }
                
                # Extract text from all pages
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text()
                    
                    page_info = {
                        'page_number': i + 1,
                        'text': page_text
                    }
                    
                    document_info['pages'].append(page_info)
                    document_info['full_text'] += page_text + "\n\n"
                
                # Create text chunks for LLM processing
                document_info['chunks'] = self._create_chunks(document_info['full_text'])
                
                return document_info
                
        except Exception as e:
            raise RuntimeError(f"Error parsing PDF: {str(e)}")
    
    def _parse_docx(self, docx_path: str) -> Dict:
        """
        Extract text and structure from DOCX documents.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Dict containing document text and metadata
        """
        document_info = {
            'file_path': docx_path,
            'file_type': 'docx',
            'sections': [],
            'full_text': "",
            'metadata': {}
        }
        
        try:
            doc = Document(docx_path)
            
            # Extract document properties
            core_properties = doc.core_properties
            document_info['metadata'] = {
                'title': core_properties.title or '',
                'author': core_properties.author or '',
                'subject': core_properties.subject or '',
                'created': str(core_properties.created) if core_properties.created else '',
                'modified': str(core_properties.modified) if core_properties.modified else ''
            }
            
            # Extract text from paragraphs
            all_text = []
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)
            
            document_info['full_text'] = "\n".join(all_text)
            
            # Extract tables if configured
            if self.config.preserve_tables:
                tables_text = []
                for i, table in enumerate(doc.tables):
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text for cell in row.cells]
                        table_rows.append(" | ".join(row_cells))
                    
                    tables_text.append(f"[TABLE {i+1}]:\n" + "\n".join(table_rows))
                
                document_info['tables'] = tables_text
                document_info['full_text'] += "\n\n" + "\n\n".join(tables_text)
            
            # Create text chunks for LLM processing
            document_info['chunks'] = self._create_chunks(document_info['full_text'])
            
            return document_info
                
        except Exception as e:
            raise RuntimeError(f"Error parsing DOCX: {str(e)}")
    
    def _create_chunks(self, text: str) -> List[str]:
        """
        Split text into chunks for LLM processing.
        
        Args:
            text: Full document text
            
        Returns:
            List of text chunks
        """
        chunk_size = self.config.chunk_size
        overlap = self.config.overlap
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Get chunk of text
            end = start + chunk_size
            if end >= len(text):
                chunk = text[start:]
            else:
                # Find a good breaking point (prefer paragraph breaks)
                # Look for newlines in the overlap region
                search_region_start = max(end - overlap, start)
                search_region = text[search_region_start:end]
                
                # Try to find double newline (paragraph break)
                break_pos = search_region.rfind('\n\n')
                if break_pos == -1:
                    # If no paragraph break, try a single newline
                    break_pos = search_region.rfind('\n')
                
                if break_pos != -1:
                    # Found a good breaking point
                    end = search_region_start + break_pos + 1
                # Else, just break at the chunk size
            
            chunk = text[start:end].strip()
            chunks.append(chunk)
            
            # Move start position, accounting for overlap
            start = end
        
        return chunks